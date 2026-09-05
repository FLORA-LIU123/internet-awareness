"""
安全评估报告生成页面
一键生成包含封面、评分摘要、等保合规、TLS详情、根因分析的Word报告。
"""
import io
from datetime import datetime, timedelta, timezone

import pandas as pd
import streamlit as st

from src.storage import db
from src.utils.config_loader import Config
from src.analysis import compliance
from app.pages._insights import _load_tls_detail, _load_health_history, _load_fused_history
from app.styles import C_HEALTHY, C_WARNING, C_CRITICAL, C_BLUE, C_TEAL, BG_CARD, BORDER, TEXT_DIM, TEXT_MAIN, rgba


def _gen_report_docx(targets_data: list, report_time: str) -> bytes:
    """生成Word格式安全评估报告，返回字节流。"""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # A4页面设置
    sec = doc.sections[0]
    sec.page_width  = Cm(21)
    sec.page_height = Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(2.8)
    sec.top_margin  = sec.bottom_margin = Cm(2.5)

    def h1(text):
        p = doc.add_heading(text, level=1)
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)
        return p

    def h2(text):
        p = doc.add_heading(text, level=2)
        for r in p.runs:
            r.font.color.rgb = RGBColor(0x1e, 0x40, 0xaf)
        return p

    def body(text):
        return doc.add_paragraph(text)

    # ── 封面 ──────────────────────────────────────────────────────────────────
    cover = doc.add_heading("", 0)
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cover.add_run("网络安全态势感知平台")
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1a, 0x56, 0xdb)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("网站安全评估报告").font.size = Pt(16)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"报告生成时间：{report_time}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)

    doc.add_paragraph()

    # ── 1. 执行摘要 ──────────────────────────────────────────────────────────
    h1("一、执行摘要")
    scores = [t["health_score"] for t in targets_data if t["health_score"] is not None]
    avg = sum(scores) / len(scores) if scores else 0
    critical = sum(1 for s in scores if s < 60)
    warning  = sum(1 for s in scores if 60 <= s < 80)
    healthy  = sum(1 for s in scores if s >= 80)

    body(
        f"本报告基于网络安全态势感知平台对 {len(targets_data)} 个目标网站的实测数据生成，"
        f"采集时间范围为报告生成前24小时。"
    )
    body(
        f"综合评估结论：平均健康度 {avg:.1f} 分，"
        f"其中正常 {healthy} 个、警告 {warning} 个、严重异常 {critical} 个。"
    )
    if critical > 0:
        p = doc.add_paragraph()
        p.add_run("⚠ 重要提示：").bold = True
        p.add_run(f"存在 {critical} 个目标健康度低于60分，建议优先处置。")

    # ── 2. 各目标安全评分概览 ─────────────────────────────────────────────────
    h1("二、各目标安全评分概览")
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, h in enumerate(["目标名称", "综合健康度", "可用性", "响应时延", "安全评分", "状态"]):
        hdr[i].text = h
        for para in hdr[i].paragraphs:
            for run in para.runs:
                run.bold = True

    for t in targets_data:
        row = tbl.add_row().cells
        row[0].text = t["name"]
        row[1].text = f"{t['health_score']:.1f}" if t["health_score"] else "N/A"
        row[2].text = f"{t['avail']:.1f}" if t["avail"] else "N/A"
        row[3].text = f"{t['resp']:.1f}" if t["resp"] else "N/A"
        row[4].text = f"{t['sec']:.1f}" if t["sec"] else "N/A"
        hs = t["health_score"] or 0
        row[5].text = "正常" if hs >= 80 else ("警告" if hs >= 60 else "严重")

    doc.add_paragraph()

    # ── 3. TLS/HTTPS安全详情 ──────────────────────────────────────────────────
    h1("三、TLS/HTTPS 安全检测详情")
    _HDR_NAMES = {
        "tls_hdr_hsts":                   "HSTS",
        "tls_hdr_csp":                    "CSP",
        "tls_hdr_x_frame_options":        "X-Frame-Options",
        "tls_hdr_x_content_type_options": "X-Content-Type-Options",
        "tls_hdr_referrer_policy":        "Referrer-Policy",
        "tls_hdr_permissions_policy":     "Permissions-Policy",
    }
    for t in targets_data:
        h2(t["name"])
        tls = t.get("tls_detail", {})
        if not tls:
            body("TLS数据尚未采集，请等待下次巡检后重新生成报告。")
            continue

        lines = [
            f"TLS综合评分：{tls.get('tls_security', 'N/A')}",
            f"证书剩余天数：{tls.get('tls_cert_days', 'N/A')} 天",
            f"TLS协议版本评分：{tls.get('tls_version_score', 'N/A')}",
            f"HTTPS强制重定向：{'是' if tls.get('tls_https_redirect', 0) else '否'}",
            f"证书透明度(SCT)：{'已配置' if tls.get('tls_sct', 0) else '未配置'}",
        ]
        for line in lines:
            doc.add_paragraph(line, style="List Bullet")

        missing_hdrs = [name for key, name in _HDR_NAMES.items() if tls.get(key, 1) == 0]
        if missing_hdrs:
            p = doc.add_paragraph()
            p.add_run("缺失安全响应头：").bold = True
            p.add_run("、".join(missing_hdrs))
        else:
            body("所有主要安全响应头均已配置。")

    # ── 4. 等保2.0合规自查 ────────────────────────────────────────────────────
    h1("四、等保2.0合规自查结果（GB/T 22239-2019）")
    body("以下结果基于实测TLS/HTTPS数据自动映射等保2.0三级技术要求条款，供参考使用。")

    for t in targets_data:
        h2(t["name"])
        tls = t.get("tls_detail", {})
        result = compliance.evaluate(tls)
        body(f"合规得分：{result['score']:.0f} 分（{result['level']}）")
        body(f"达标 {result['pass_count']} 项 / 需关注 {result['warn_count']} 项 / 不达标 {result['fail_count']} 项")

        # 只列出不达标和需关注的条款
        issues = [c for c in result["clauses"] if c["status"] in ("fail", "warn")]
        if issues:
            p = doc.add_paragraph()
            p.add_run("需整改条款：").bold = True
            for clause in issues:
                status_label = "不达标" if clause["status"] == "fail" else "需关注"
                doc.add_paragraph(
                    f"[{status_label}] {clause['id']} {clause['name']}：{clause['detail']}。"
                    f"修复建议：{clause['fix']}",
                    style="List Bullet"
                )
        else:
            body("所有检测条款均已达标。")

    # ── 5. 根因分析摘要 ───────────────────────────────────────────────────────
    h1("五、评分骤降事件根因分析")
    from src.analysis import root_cause
    cfg = Config.get()
    any_event = False
    for t in targets_data:
        events = root_cause.analyze(cfg.db_path, t["name"], hours=24)
        if not events:
            continue
        any_event = True
        h2(t["name"])
        for ev in events:
            t_str = ev["time"].strftime("%Y-%m-%d %H:%M") if hasattr(ev["time"], "strftime") else str(ev["time"])
            p = doc.add_paragraph()
            p.add_run(f"{t_str}  评分 {ev['from_score']:.1f} → {ev['to_score']:.1f}（下降{ev['drop']:.1f}分）").bold = True
            doc.add_paragraph(ev["summary"], style="List Bullet")

    if not any_event:
        body("过去24小时内未检测到评分骤降事件。")

    # ── 6. 改进建议汇总 ───────────────────────────────────────────────────────
    h1("六、综合改进建议")
    all_issues = []
    for t in targets_data:
        tls = t.get("tls_detail", {})
        result = compliance.evaluate(tls)
        for clause in result["clauses"]:
            if clause["status"] in ("fail", "warn"):
                all_issues.append((t["name"], clause))

    if all_issues:
        body(f"共发现 {len(all_issues)} 项需整改内容，按优先级排序建议如下：")
        # 按weight降序
        all_issues.sort(key=lambda x: x[1]["weight"], reverse=True)
        for name, clause in all_issues[:10]:
            doc.add_paragraph(
                f"【{name}】{clause['name']}：{clause['fix']}",
                style="List Number"
            )
        if len(all_issues) > 10:
            body(f"（另有 {len(all_issues) - 10} 项次要问题，详见各目标等保合规章节）")
    else:
        body("所有目标当前无明显需整改项。")

    # 结尾
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run(f"本报告由网络安全态势感知平台自动生成  ·  {report_time}")
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(0x9c, 0xa3, 0xaf)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def render():
    cfg = Config.get()
    db_path = cfg.db_path

    st.markdown("## 📄 安全评估报告")
    st.caption("选择监测目标，一键生成结构化安全评估报告（Word格式），包含等保合规自查、TLS详情和根因分析")
    st.markdown("---")

    # 说明卡片
    st.markdown(
        f'<div style="background:{rgba(C_BLUE,0.06)};border:1px solid {rgba(C_BLUE,0.2)};'
        f'border-radius:10px;padding:14px 18px;margin-bottom:20px;">'
        f'<div style="font-size:0.9rem;font-weight:700;color:{C_BLUE};margin-bottom:8px;">📋 报告内容包括</div>'
        f'<div style="font-size:0.83rem;color:{TEXT_MAIN};line-height:2;">'
        f'一、执行摘要（整体评级）&emsp; 二、各目标安全评分概览<br>'
        f'三、TLS/HTTPS 安全检测详情 &emsp; 四、等保2.0合规自查（GB/T 22239-2019）<br>'
        f'五、评分骤降事件根因分析 &emsp; 六、综合改进建议（按优先级排序）'
        f'</div></div>',
        unsafe_allow_html=True,
    )

    # 加载数据
    health_df = _load_health_history(db_path, hours=24)
    fused_df  = _load_fused_history(db_path, hours=24)

    if health_df.empty:
        st.warning("暂无数据，请等待首次采集完成后再生成报告。")
        return

    all_targets = sorted(health_df["target_name"].unique().tolist())

    # ── 目标选择区 ────────────────────────────────────────────────────────────
    st.markdown("### 选择报告目标")
    col_sel, col_tip = st.columns([3, 2])
    with col_sel:
        selected_targets = st.multiselect(
            "选择要纳入报告的监测目标（可多选）",
            options=all_targets,
            default=all_targets,
            key="report_targets",
        )
    with col_tip:
        st.markdown(
            f'<div style="background:{rgba(C_TEAL,0.07)};border:1px solid {rgba(C_TEAL,0.2)};'
            f'border-radius:8px;padding:10px 14px;font-size:0.8rem;color:{TEXT_DIM};margin-top:8px;">'
            f'💡 可为单个目标单独生成报告，适合向不同机构分别汇报</div>',
            unsafe_allow_html=True,
        )

    if not selected_targets:
        st.info("请至少选择一个监测目标。")
        return

    # ── 所选目标预览 ──────────────────────────────────────────────────────────
    st.markdown("### 目标当前状态")
    cols = st.columns(len(selected_targets))
    targets_data = []
    for col, name in zip(cols, selected_targets):
        sub_h = health_df[health_df["target_name"] == name]
        health_score = float(sub_h.sort_values("scored_at").iloc[-1]["score"]) if not sub_h.empty else None

        avail = resp = sec = None
        if not fused_df.empty:
            sub_f = fused_df[fused_df["target_name"] == name]
            if not sub_f.empty:
                latest = sub_f.sort_values("fused_at").iloc[-1]
                avail = float(latest.get("availability_score") or 0)
                resp  = float(latest.get("response_time_score") or 0)
                sec   = float(latest.get("security_score") or 0)

        tls_detail = _load_tls_detail(db_path, name)
        comp_result = compliance.evaluate(tls_detail)

        targets_data.append({
            "name": name,
            "health_score": health_score,
            "avail": avail,
            "resp": resp,
            "sec": sec,
            "tls_detail": tls_detail,
        })

        hs = health_score or 0
        comp_score = comp_result["score"]
        comp_color = C_HEALTHY if comp_score >= 75 else (C_WARNING if comp_score >= 50 else C_CRITICAL)

        col.markdown(
            f'<div style="background:{rgba(C_BLUE,0.05)};border:1px solid {BORDER};'
            f'border-radius:10px;padding:12px 14px;text-align:center;">'
            f'<div style="font-size:0.82rem;font-weight:700;color:{TEXT_MAIN};margin-bottom:8px;">{name}</div>'
            f'<div style="font-size:1.5rem;font-weight:800;color:{"#22c55e" if hs>=80 else ("#f59e0b" if hs>=60 else "#ef4444")};">{hs:.1f}</div>'
            f'<div style="font-size:0.7rem;color:{TEXT_DIM};">综合健康度</div>'
            f'<div style="margin-top:8px;font-size:0.75rem;color:{comp_color};">等保合规 {comp_score:.0f}分 · {comp_result["level"]}</div>'
            f'</div>',
            unsafe_allow_html=True,
        )

    st.markdown("---")

    # ── 生成按钮 ──────────────────────────────────────────────────────────────
    btn_label = f"🚀 生成报告（{len(selected_targets)} 个目标）"
    if st.button(btn_label, type="primary", use_container_width=True):
        with st.spinner("正在生成报告，请稍候..."):
            report_time = datetime.now().strftime("%Y年%m月%d日 %H:%M")
            try:
                docx_bytes = _gen_report_docx(targets_data, report_time)
                # 文件名包含目标名（单目标时更直观）
                if len(selected_targets) == 1:
                    safe_name = selected_targets[0].replace("/", "_").replace("\\", "_")[:20]
                    filename = f"安全评估报告_{safe_name}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                else:
                    filename = f"网络安全评估报告_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.success(f"报告生成成功！共 {len(selected_targets)} 个目标，点击下方按钮下载。")
                st.download_button(
                    label="⬇ 下载 Word 报告",
                    data=docx_bytes,
                    file_name=filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except ImportError:
                st.error("缺少 python-docx 依赖，请在终端执行：.venv\\Scripts\\pip install python-docx")
            except Exception as e:
                st.error(f"报告生成失败：{e}")